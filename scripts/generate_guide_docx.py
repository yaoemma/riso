from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parents[1] / "Guide-prise-en-main-et-deploiement-RISO.docx"


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    shading = p.paragraph_format
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph("")


def build():
    doc = Document()

    add_title(doc, "Guide de prise en main et de déploiement")
    sub = doc.add_paragraph(
        "Site web RISO — React · JavaScript · CSS · Vite"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_paragraph("")
    sub2 = doc.add_paragraph("Réseau Ivoirien des Spécialistes de l'Orientation")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(11)

    doc.add_paragraph("")
    meta = doc.add_paragraph("Document généré pour l'équipe RISO")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "Ce document explique comment installer le site en local, modifier son contenu au quotidien "
        "et le publier en ligne (déploiement). Le site RISO est une application web construite avec "
        "React, JavaScript, CSS et Vite. Tout le code source se trouve dans le dossier riso-react/."
    )

    add_heading(doc, "1.1. Stack technique : React, JavaScript, CSS et Vite", level=2)
    add_table(
        doc,
        ["Technologie", "Rôle dans le site RISO", "Fichiers concernés"],
        [
            [
                "React",
                "Construit l'interface (pages, menu, formulaires, annuaire écoles). "
                "Les composants sont réutilisables et se mettent à jour sans recharger toute la page.",
                "src/pages/, src/components/, src/App.jsx, src/main.jsx",
            ],
            [
                "JavaScript",
                "Langage de programmation du projet. Gère la logique : navigation, forum, "
                "filtres de l'annuaire, carrousel, données des écoles.",
                "src/**/*.jsx, src/**/*.js, src/hooks/, src/utils/, src/data/",
            ],
            [
                "CSS",
                "Définit l'apparence : couleurs RISO (bleu, orange), mise en page, menu, "
                "en-têtes de pages, bandeau Flash info, version mobile.",
                "src/styles/main.css, src/styles/forum.css",
            ],
            [
                "Vite",
                "Outil de développement et de compilation. Lance le site en local (npm run dev) "
                "et produit les fichiers finaux pour la mise en ligne (npm run build → dist/).",
                "vite.config.js, package.json",
            ],
        ],
    )
    add_para(doc, "Compléments utilisés avec cette stack :", bold=True)
    add_bullet(doc, "React Router — gestion des URLs et navigation entre les pages")
    add_bullet(doc, "JSX — syntaxe mélangeant JavaScript et HTML dans les fichiers .jsx")
    add_bullet(doc, "Aucune base de données ni fichier .env requis pour faire fonctionner le site")

    add_heading(doc, "1.2. Comment ces technologies travaillent ensemble", level=2)
    add_number(doc, "Vite démarre le serveur local et compile le projet.")
    add_number(doc, "React affiche les composants (Navbar, pages, Footer…) dans le navigateur.")
    add_number(doc, "JavaScript alimente la logique métier (données écoles, forum, interactions).")
    add_number(doc, "CSS applique le design visuel par-dessus la structure React.")
    add_number(doc, "Au déploiement, Vite génère des fichiers HTML/CSS/JS optimisés dans dist/.")

    add_heading(doc, "2. Prérequis")
    add_para(doc, "Avant de commencer, installez les outils suivants sur votre ordinateur :")
    add_table(
        doc,
        ["Outil", "Version minimale", "Vérification"],
        [
            ["Node.js", "20 ou plus récent", "node -v"],
            ["npm", "Inclus avec Node.js", "npm -v"],
            ["Git", "Recommandé", "git --version"],
            ["Éditeur de code", "Cursor, VS Code…", "—"],
        ],
    )
    add_para(doc, "Téléchargement Node.js : https://nodejs.org/")
    add_para(
        doc,
        "Connaissance recommandée : notions de HTML, CSS et JavaScript, navigation dans les dossiers, "
        "utilisation d'un terminal (PowerShell sous Windows). Une initiation à React (.jsx) est un plus "
        "mais non obligatoire pour les modifications de contenu simples."
    )

    add_heading(doc, "3. Structure du projet")
    add_code(
        doc,
        "projet_RISO1/\n"
        "├── README.md\n"
        "├── .gitignore\n"
        "└── riso-react/\n"
        "    ├── public/images/          # Images du site (logos, événements, écoles…)\n"
        "    ├── src/\n"
        "    │   ├── App.jsx             # Définition des routes (URLs)\n"
        "    │   ├── main.jsx            # Point d'entrée React\n"
        "    │   ├── pages/              # Une page par section du site\n"
        "    │   ├── components/         # Blocs réutilisables (menu, pied de page…)\n"
        "    │   ├── data/               # Données écoles, universités, forum\n"
        "    │   ├── hooks/              # Logique React (carrousel, menu mobile…)\n"
        "    │   ├── utils/              # Fonctions utilitaires\n"
        "    │   └── styles/             # Fichiers CSS\n"
        "    ├── package.json            # Dépendances et scripts npm\n"
        "    └── vite.config.js          # Configuration Vite"
    )

    add_heading(doc, "4. Installation et premier lancement")
    add_heading(doc, "4.1. Récupérer le projet", level=2)
    add_number(doc, "Ouvrir un terminal (PowerShell).")
    add_number(doc, "Se placer dans le dossier souhaité, puis cloner le dépôt Git (ou copier le dossier du projet).")
    add_number(doc, "Entrer dans le dossier riso-react :")
    add_code(doc, "cd \"chemin\\vers\\projet_RISO1\\riso-react\"")

    add_heading(doc, "4.2. Installer les dépendances", level=2)
    add_code(doc, "npm install")
    add_para(doc, "Cette commande télécharge une seule fois les bibliothèques nécessaires (dossier node_modules/).")

    add_heading(doc, "4.3. Lancer le site en local", level=2)
    add_code(doc, "npm run dev")
    add_para(doc, "Le terminal affiche une adresse locale, en général : http://localhost:5173")
    add_para(doc, "Ouvrez cette adresse dans votre navigateur. Le site se recharge automatiquement à chaque modification de fichier.")
    add_para(doc, "Pour arrêter le serveur : Ctrl + C dans le terminal.")

    add_heading(doc, "4.4. Vérifier le build de production", level=2)
    add_para(doc, "Avant chaque déploiement, vérifiez que le site se compile sans erreur :")
    add_code(doc, "npm run build\nnpm run preview")
    add_para(doc, "preview ouvre la version compilée (souvent sur http://localhost:4173). Si build échoue, corrigez les erreurs avant de déployer.")

    add_heading(doc, "5. Prise en main — modifier le contenu")
    add_heading(doc, "5.1. Messages Flash info (bandeau du haut)", level=2)
    add_para(doc, "Fichier : riso-react/src/components/layout/InfoBar.jsx")
    add_para(doc, "Le bandeau défilant affiche les actualités importantes (rentrée scolaire, affectation 6ème, orientation Seconde…).")
    add_para(doc, "Pour modifier les messages :", bold=True)
    add_number(doc, "Ouvrir InfoBar.jsx.")
    add_number(doc, "Modifier le tableau FLASH_MESSAGES.")
    add_number(doc, "Pour un message simple, utiliser une chaîne de texte entre guillemets.")
    add_number(doc, "Pour un message avec lien vers mendob-ci.org, utiliser le format objet { before, linkLabel, after }.")
    add_number(doc, "Enregistrer le fichier : le site local se met à jour automatiquement.")
    add_para(doc, "Vitesse du défilement : riso-react/src/styles/main.css, règle .flash-scroll (animation scroll-left, durée en secondes).")

    add_heading(doc, "5.2. Menu de navigation et liens sociaux", level=2)
    add_para(doc, "Menu principal : riso-react/src/components/layout/Navbar.jsx")
    add_para(doc, "Liens Facebook et WhatsApp : InfoBar.jsx (section social-media) et Footer.jsx")

    add_heading(doc, "5.3. Pages du site", level=2)
    add_table(
        doc,
        ["URL", "Fichier page", "Contenu principal"],
        [
            ["/", "pages/Home.jsx", "Page d'accueil"],
            ["/apropos", "pages/AproposPage.jsx", "À propos du RISO"],
            ["/membres", "pages/MembresPage.jsx", "Spécialistes"],
            ["/histoire", "pages/HistoirePage.jsx", "Histoire"],
            ["/valeurs", "pages/ValeursPage.jsx", "Valeurs"],
            ["/activites", "pages/ActivitesPage.jsx", "Activités"],
            ["/nos-services", "pages/NosServicesPage.jsx", "Services"],
            ["/actualites", "pages/ActualitesPage.jsx", "Actualités et médias"],
            ["/contact", "pages/ContactPage.jsx", "Contact"],
            ["/ecoles-universites", "pages/EcolesUniversitesPage.jsx", "Annuaire établissements"],
            ["/temoignages", "pages/TemoignagesPage.jsx", "Témoignages"],
            ["/forum", "pages/ForumEchange.jsx", "Forum d'échange"],
            ["/assemblee-generale", "pages/AssembleeGenerale.jsx", "Article AG"],
            ["/formation-esatic", "pages/FormationEsatic.jsx", "Article ESATIC"],
            ["/forum-emploi", "pages/ForumEmploi.jsx", "Article forum emploi"],
            ["/formation-ena", "pages/FormationEna.jsx", "Article formation ENA"],
        ],
    )
    add_para(
        doc,
        "Chaque page assemble un en-tête (PageHeader ou ArticlePageHeader) et un ou plusieurs composants "
        "dans src/components/home/."
    )

    add_heading(doc, "5.4. En-têtes visuels des pages", level=2)
    add_para(doc, "Configuration des images d'en-tête : riso-react/src/utils/pageHeaderVisuals.js")
    add_para(doc, "Styles CSS des en-têtes : riso-react/src/styles/main.css (section .page-header)")

    add_heading(doc, "5.5. Annuaire Écoles & Universités", level=2)
    add_para(doc, "Page : /ecoles-universites — composant : components/home/SchoolsPartners.jsx")
    add_table(
        doc,
        ["Fichier de données", "Rôle"],
        [
            ["data/ivorianSchools.js", "Fusion des listes et affichage grille / Voir plus"],
            ["data/excellenceSchools.js", "Établissements publics d'excellence"],
            ["data/primarySchools.js", "Écoles primaires privées et catholiques"],
            ["data/secondarySchools.js", "Collèges et lycées privés / catholiques"],
            ["data/menaSecondarySchools.json", "Collèges publics (source MENA)"],
            ["data/ivorianInstitutions.js", "Universités et grandes écoles"],
            ["utils/schoolVisuals.js", "Couleurs et initiales des cartes sans logo"],
        ],
    )
    add_para(doc, "Logos universités : public/images/schools/logos/")

    add_heading(doc, "5.6. Images et médias", level=2)
    add_para(doc, "Toutes les images statiques sont dans riso-react/public/images/.")
    add_para(doc, "Dans le code, on y accède avec un chemin commençant par /images/ (ex. /images/logo/logo-riso-navbar-112.png).")
    add_para(doc, "Après ajout d'une nouvelle image, redémarrer npm run dev si elle n'apparaît pas.")

    add_heading(doc, "5.7. Forum d'échange", level=2)
    add_para(doc, "Page : pages/ForumEchange.jsx — logique : hooks/useForum.js — styles : styles/forum.css")
    add_para(doc, "Les messages du forum sont stockés localement dans le navigateur (localStorage), pas sur un serveur.")

    add_heading(doc, "5.8. Styles visuels (couleurs, mise en page)", level=2)
    add_para(doc, "Fichier principal : riso-react/src/styles/main.css")
    add_para(doc, "Couleurs RISO courantes : bleu #2E78C0, orange #F28B2E.")

    add_heading(doc, "6. Commandes npm utiles")
    add_table(
        doc,
        ["Commande", "Description"],
        [
            ["npm run dev", "Serveur de développement local"],
            ["npm run build", "Compile le site pour la production (dossier dist/)"],
            ["npm run preview", "Prévisualise le build de production"],
            ["npm run lint", "Vérifie la qualité du code (Oxlint)"],
        ],
    )

    add_heading(doc, "7. Déploiement — mettre le site en ligne")
    add_para(
        doc,
        "Le site RISO est une Single Page Application (SPA) : après compilation, il produit des fichiers "
        "statiques HTML, CSS et JavaScript dans le dossier dist/. Il faut un hébergeur capable de "
        "servir ces fichiers et de rediriger toutes les URLs vers index.html."
    )

    add_heading(doc, "7.1. Étape obligatoire : build", level=2)
    add_code(doc, "cd riso-react\nnpm run build")
    add_para(doc, "Le dossier riso-react/dist/ contient le site prêt à publier.")

    add_heading(doc, "7.2. Déploiement sur Netlify (recommandé)", level=2)
    add_para(doc, "Netlify convient bien aux sites React/Vite et gère automatiquement les routes SPA.")
    add_para(doc, "Méthode A — Depuis Git (recommandée) :", bold=True)
    add_number(doc, "Créer un compte sur https://www.netlify.com/")
    add_number(doc, "Connecter le dépôt GitHub/GitLab/Bitbucket du projet.")
    add_number(doc, "Configurer le build :")
    add_bullet(doc, "Base directory : riso-react")
    add_bullet(doc, "Build command : npm run build")
    add_bullet(doc, "Publish directory : riso-react/dist")
    add_number(doc, "Créer un fichier netlify.toml à la racine du dépôt avec le contenu suivant :")
    add_code(
        doc,
        "[build]\n"
        "  base = \"riso-react\"\n"
        "  command = \"npm run build\"\n"
        "  publish = \"riso-react/dist\"\n\n"
        "[[redirects]]\n"
        "  from = \"/*\"\n"
        "  to = \"/index.html\"\n"
        "  status = 200"
    )
    add_number(doc, "Pousser les modifications sur Git : Netlify rebuild et publie automatiquement.")
    add_para(doc, "Méthode B — Glisser-déposer (sans Git) :", bold=True)
    add_number(doc, "Exécuter npm run build en local.")
    add_number(doc, "Sur Netlify : Sites → Add new site → Deploy manually.")
    add_number(doc, "Glisser le dossier dist/ dans la zone de déploiement.")

    add_heading(doc, "7.3. Déploiement sur un hébergeur classique (o2switch, LWS, VPS…)", level=2)
    add_number(doc, "Exécuter npm run build en local.")
    add_number(doc, "Envoyer tout le contenu de dist/ sur le serveur (FTP, SFTP ou gestionnaire de fichiers).")
    add_number(doc, "Placer les fichiers à la racine du site web (public_html ou www).")
    add_number(doc, "Configurer la réécriture d'URL pour que toutes les routes pointent vers index.html.")
    add_para(doc, "Exemple Apache (.htaccess à la racine) :")
    add_code(
        doc,
        "<IfModule mod_rewrite.c>\n"
        "  RewriteEngine On\n"
        "  RewriteBase /\n"
        "  RewriteRule ^index\\.html$ - [L]\n"
        "  RewriteCond %{REQUEST_FILENAME} !-f\n"
        "  RewriteCond %{REQUEST_FILENAME} !-d\n"
        "  RewriteRule . /index.html [L]\n"
        "</IfModule>"
    )

    add_heading(doc, "7.4. Déploiement sur GitHub Pages", level=2)
    add_number(doc, "Adapter vite.config.js pour ajouter base: '/nom-du-repo/' si le site n'est pas à la racine du domaine.")
    add_number(doc, "npm run build")
    add_number(doc, "Publier le contenu de dist/ sur la branche gh-pages ou via GitHub Actions.")

    add_heading(doc, "7.5. Vérifications après déploiement", level=2)
    add_bullet(doc, "La page d'accueil s'affiche correctement.")
    add_bullet(doc, "Tester plusieurs pages : /apropos, /ecoles-universites, /forum…")
    add_bullet(doc, "Rafraîchir (F5) une page interne : elle ne doit pas afficher d'erreur 404.")
    add_bullet(doc, "Vérifier les images, le menu mobile et les liens externes (mendob-ci.org, Facebook, WhatsApp).")
    add_bullet(doc, "Tester sur téléphone (responsive).")

    add_heading(doc, "8. Workflow de mise à jour")
    add_number(doc, "Modifier les fichiers en local.")
    add_number(doc, "Vérifier avec npm run dev.")
    add_number(doc, "Lancer npm run build pour confirmer qu'il n'y a pas d'erreur.")
    add_number(doc, "Commiter et pousser sur Git (si déploiement automatique) ou redéployer dist/.")
    add_para(doc, "Bonnes pratiques :", bold=True)
    add_bullet(doc, "Ne jamais modifier directement le dossier dist/ (il est régénéré à chaque build).")
    add_bullet(doc, "Ne pas envoyer node_modules/ sur le serveur ni sur Git.")
    add_bullet(doc, "Sauvegarder les images avant de les remplacer.")

    add_heading(doc, "9. Dépannage")
    add_table(
        doc,
        ["Problème", "Solution"],
        [
            ["npm install échoue", "Vérifier Node.js 20+, supprimer node_modules/ et relancer npm install"],
            ["Page blanche après déploiement", "Vérifier la redirection SPA vers index.html"],
            ["404 en rafraîchissant une page", "Configurer redirects Netlify ou .htaccess"],
            ["Images absentes", "Vérifier le chemin /images/... et que le fichier est dans public/images/"],
            ["npm run build échoue", "Lire le message d'erreur dans le terminal et corriger le fichier indiqué"],
            ["Port 5173 déjà utilisé", "Fermer l'autre serveur ou utiliser le port proposé par Vite"],
        ],
    )

    add_heading(doc, "10. Contacts et ressources")
    add_para(doc, "Documentation des technologies du projet :", bold=True)
    add_bullet(doc, "React : https://react.dev/")
    add_bullet(doc, "JavaScript (MDN) : https://developer.mozilla.org/fr/docs/Web/JavaScript")
    add_bullet(doc, "CSS (MDN) : https://developer.mozilla.org/fr/docs/Web/CSS")
    add_bullet(doc, "Vite : https://vite.dev/")
    add_para(doc, "Autres ressources :", bold=True)
    add_bullet(doc, "Portail orientation MENA : https://mendob-ci.org/")
    add_bullet(doc, "Documentation Netlify : https://docs.netlify.com/")

    doc.save(OUTPUT)
    print(f"Document créé : {OUTPUT}")


if __name__ == "__main__":
    build()
